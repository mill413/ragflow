#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from io import BytesIO

from docx import Document

from deepdoc.parser.docx_parser import RAGFlowDocxParser


def test_docx_parser_extracts_nested_table_content():
    document = Document()
    outer = document.add_table(rows=2, cols=2)
    outer.cell(0, 0).text = "outer header"
    outer.cell(0, 1).text = "outer value"
    outer.cell(1, 1).text = "outer row value"

    nested = outer.cell(1, 0).add_table(rows=2, cols=2)
    nested.cell(0, 0).text = "nested A"
    nested.cell(0, 1).text = "nested B"
    nested.cell(1, 0).text = "nested C"
    nested.cell(1, 1).text = "nested D"

    output = BytesIO()
    document.save(output)

    _, tables = RAGFlowDocxParser()(output.getvalue())
    extracted = "\n".join(
        line for table in tables for line in table if isinstance(line, str)
    )

    assert "nested A" in extracted
    assert "nested B" in extracted
    assert "nested C" in extracted
    assert "nested D" in extracted


def test_docx_parser_extracts_multiple_nested_table_levels():
    document = Document()
    outer = document.add_table(rows=2, cols=1)
    outer.cell(0, 0).text = "header"
    middle = outer.cell(1, 0).add_table(rows=1, cols=1)
    inner = middle.cell(0, 0).add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "deeply nested value"

    output = BytesIO()
    document.save(output)

    _, tables = RAGFlowDocxParser()(output.getvalue())
    extracted = "\n".join(
        line for table in tables for line in table if isinstance(line, str)
    )

    assert "deeply nested value" in extracted


def test_naive_docx_parser_preserves_nested_table_html():
    from rag.app.naive import Docx as NaiveDocxParser

    document = Document()
    outer = document.add_table(rows=1, cols=2)
    nested = outer.cell(0, 0).add_table(rows=2, cols=2)
    nested.cell(0, 0).text = "nested A"
    nested.cell(0, 1).text = "nested B"
    nested.cell(1, 0).text = "nested C"
    nested.cell(1, 1).text = "nested D"
    outer.cell(0, 1).text = "outer value"

    output = BytesIO()
    document.save(output)

    sections = NaiveDocxParser()("nested.docx", binary=output.getvalue())
    html = "\n".join(table or "" for _, _, table in sections)

    assert html.count("<table>") == 2
    assert "<td>nested A</td>" in html
    assert "<td>nested B</td>" in html
    assert "<td>nested C</td>" in html
    assert "<td>nested D</td>" in html
    assert "<td>outer value</td>" in html
