#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import gc
import importlib.util
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile


def _load_large_ooxml_module():
    module_path = (
        Path(__file__).parents[4] / "deepdoc" / "parser" / "large_ooxml.py"
    )
    spec = importlib.util.spec_from_file_location("large_ooxml", module_path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def test_large_ooxml_text_nodes_are_supported():
    module = _load_large_ooxml_module()
    module.configure_large_ooxml_parsers()

    from docx.oxml import parse_xml as parse_docx_xml
    from openpyxl.xml.functions import fromstring as parse_excel_xml
    from pptx.oxml import parse_xml as parse_pptx_xml

    text_size = 11 * 1024 * 1024
    large_text = "x" * text_size
    docx_root = parse_docx_xml(
        f'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:r><w:t>{large_text}</w:t></w:r></w:p></w:body></w:document>'
    )
    assert len(docx_root.xpath("string(.//w:t)")) == text_size
    del docx_root, large_text
    gc.collect()

    large_text = "x" * text_size
    excel_root = parse_excel_xml(f"<worksheet><value>{large_text}</value></worksheet>")
    assert len(excel_root.find("value").text) == text_size
    del excel_root, large_text
    gc.collect()

    large_text = "x" * text_size
    pptx_root = parse_pptx_xml(
        f'<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:extLst>{large_text}</p:extLst></p:presentation>'
    )
    assert len(pptx_root.xpath("string(.//p:extLst)")) == text_size


def test_docm_opens_without_rewriting_or_executing_macro_parts():
    module = _load_large_ooxml_module()
    from docx import Document
    module.configure_large_ooxml_parsers()

    document = Document()
    document.add_paragraph("DOCM content")
    docx = BytesIO()
    document.save(docx)

    docm = BytesIO()
    with ZipFile(BytesIO(docx.getvalue())) as source, ZipFile(
        docm, "w", ZIP_DEFLATED
    ) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "[Content_Types].xml":
                data = data.replace(
                    b"application/vnd.openxmlformats-officedocument."
                    b"wordprocessingml.document.main+xml",
                    b"application/vnd.ms-word.document.macroEnabled.main+xml",
                )
            target.writestr(info, data)
        target.writestr("word/vbaProject.bin", b"macro payload")

    original = docm.getvalue()
    parsed = Document(BytesIO(original))

    assert [paragraph.text for paragraph in parsed.paragraphs] == [
        "DOCM content"
    ]
    with ZipFile(BytesIO(original)) as archive:
        assert archive.read("word/vbaProject.bin") == b"macro payload"
